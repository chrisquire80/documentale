import asyncio
import os
import sys

from docx import Document

from app.services.ocr import extract_text

async def test_docx():
    print("Creazione file docx di test...")
    doc = Document()
    doc.add_heading('Documento di Test', 0)
    doc.add_paragraph('Questo è un testo di prova per verificare l\'estrazione OCR da file Word.')
    test_path = "/tmp/test_ocr.docx"
    doc.save(test_path)

    print(f"File salvato in {test_path}")
    print("Avvio estrazione...")

    try:
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        extracted = await extract_text(test_path, mime_type)
        print("--- TESTO ESTRATTO ---")
        print(extracted)
        print("----------------------")
        
        if "Questo è un testo di prova" in extracted:
            print("SUCCESSO!")
            sys.exit(0)
        else:
            print("ERRORE: Testo non trovato.")
            sys.exit(1)
            
    finally:
        if os.path.exists(test_path):
            os.remove(test_path)

if __name__ == "__main__":
    asyncio.run(test_docx())
